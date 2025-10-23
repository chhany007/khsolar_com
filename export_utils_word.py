"""
Word Document Export Helper Functions
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
from typing import List
from models import SimulationResult, Device, FinancialAnalysis
from report_translations import REPORT_LABELS as RL

def set_cell_background(cell, color):
    """Set background color for table cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_professional_header(doc):
    """Add KHSolar professional header"""
    # Company name
    header = doc.add_heading(RL['company_name'], 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header.runs[0]
    header_run.font.size = Pt(28)
    header_run.font.color.rgb = RGBColor(255, 107, 53)  # Orange
    header_run.font.bold = True
    
    # Subtitle
    subtitle = doc.add_paragraph(RL['company_subtitle'])
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor(44, 62, 80)
    
    doc.add_paragraph()
    
    # Report title
    title = doc.add_heading(RL['report_title'], 1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(44, 62, 80)
    
    # Date
    date_para = doc.add_paragraph(f"{RL['report_generated']}: {datetime.datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.runs[0]
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(127, 140, 141)
    
    doc.add_paragraph()

def add_customer_info(doc, system_config):
    """Add customer information section"""
    if system_config.get('customer_name'):
        doc.add_heading(RL['customer_info'], 2)
        
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Data
        data = [
            (RL['customer_name'] + ':', system_config.get('customer_name', 'N/A')),
            (RL['company'] + ':', system_config.get('customer_company', 'N/A') or RL['individual']),
            (RL['phone'] + ':', system_config.get('customer_phone', 'N/A')),
            (RL['email'] + ':', system_config.get('customer_email', 'N/A') or 'N/A'),
            (RL['address'] + ':', system_config.get('customer_address', 'N/A') or 'N/A'),
        ]
        
        for i, (label, value) in enumerate(data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value
            # Bold labels
            table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
            set_cell_background(table.rows[i].cells[0], 'ECF0F1')
        
        doc.add_paragraph()

def add_system_config(doc, system_config, device_count):
    """Add system configuration section"""
    doc.add_heading(RL['system_config'], 2)
    
    table = doc.add_table(rows=10, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Header
    header_cells = table.rows[0].cells
    header_cells[0].text = RL['parameter']
    header_cells[1].text = RL['value']
    header_cells[2].text = RL['unit']
    
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, '3498DB')
    
    # Data
    data = [
        (RL['location'], system_config.get('location', 'Cambodia'), ''),
        (RL['pv_capacity'], f"{system_config.get('pv_capacity', 0):.2f}", 'kW'),
        (RL['battery_capacity'], f"{system_config.get('battery_capacity', 0):.2f}", 'kWh'),
        (RL['inverter_power'], f"{system_config.get('inverter_power', 0):.2f}", 'kW'),
        (RL['num_devices'], str(device_count), 'units'),
        (RL['system_type'], RL['grid_tied_battery'], ''),
        (RL['labor_cost'], f"${system_config.get('labor_cost', 0):.2f}", ''),
        (RL['support_materials'], f"${system_config.get('support_material_cost', 0):.2f}", ''),
    ]
    
    for i, (param, value, unit) in enumerate(data, 1):
        table.rows[i].cells[0].text = param
        table.rows[i].cells[1].text = value
        table.rows[i].cells[2].text = unit
    
    doc.add_paragraph()

def add_cost_breakdown(doc, system_config, financial):
    """Add cost breakdown section"""
    doc.add_heading(RL['cost_breakdown'], 2)
    
    equipment_cost = system_config.get('equipment_cost', 0)
    labor_cost = system_config.get('labor_cost', 0)
    support_cost = system_config.get('support_material_cost', 0)
    total = financial.total_system_cost
    
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Medium Grid 3 Accent 1'
    
    # Header
    header_cells = table.rows[0].cells
    header_cells[0].text = RL['cost_component']
    header_cells[1].text = RL['amount']
    header_cells[2].text = RL['percentage']
    
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    # Data
    inverter_kw = system_config.get('inverter_power', 0)
    labor_note = RL['system_5kw'] if inverter_kw <= 5 else RL['system_above_5kw']
    
    data = [
        (RL['equipment'], f"${equipment_cost:,.2f}", f"{(equipment_cost/total*100):.1f}%" if total > 0 else '0%'),
        (f"{RL['labor_system']} ({labor_note})", f"${labor_cost:,.2f}", f"{(labor_cost/total*100):.1f}%" if total > 0 else '0%'),
        (RL['support_installation'], f"${support_cost:,.2f}", f"{(support_cost/total*100):.1f}%" if total > 0 else '0%'),
        (RL['total_system_cost'], f"${total:,.2f}", '100%'),
    ]
    
    for i, (component, amount, percentage) in enumerate(data, 1):
        table.rows[i].cells[0].text = component
        table.rows[i].cells[1].text = amount
        table.rows[i].cells[2].text = percentage
        
        if i == len(data):  # Last row (total)
            for cell in table.rows[i].cells:
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()

def add_financial_analysis(doc, financial):
    """Add financial analysis section"""
    doc.add_heading(RL['financial_analysis'], 2)
    
    table = doc.add_table(rows=8, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Header
    header_cells = table.rows[0].cells
    header_cells[0].text = RL['metric']
    header_cells[1].text = RL['value']
    header_cells[2].text = RL['details']
    
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, '27AE60')
    
    # Data
    data = [
        (RL['total_investment'], f"${financial.total_system_cost:,.2f}", RL['complete_installation']),
        (RL['monthly_savings'], f"${financial.monthly_savings:,.2f}", RL['electricity_reduction']),
        (RL['annual_savings'], f"${financial.annual_savings:,.2f}", RL['year1_savings']),
        (RL['payback_period'], f"{financial.payback_period_years:.1f} years", RL['breakeven_point']),
        (RL['roi'], f"{financial.roi_percent:.1f}%", RL['over_25years']),
        (RL['lifetime_savings'], f"${financial.lifetime_savings:,.2f}", RL['total_25years']),
        (RL['co2_reduction'], f"{financial.co2_reduction_kg_per_year:,.0f} kg/year", RL['environmental_impact']),
    ]
    
    for i, (metric, value, detail) in enumerate(data, 1):
        table.rows[i].cells[0].text = metric
        table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        table.rows[i].cells[1].text = value
        table.rows[i].cells[2].text = detail
    
    doc.add_paragraph()

def add_energy_summary(doc, simulation_results):
    """Add energy summary section"""
    doc.add_heading(RL['energy_analysis'], 2)
    
    total_pv = sum(r.pv_generation_kw for r in simulation_results)
    total_load = sum(r.load_kw for r in simulation_results)
    total_grid_import = sum(r.grid_import_kw for r in simulation_results)
    total_grid_export = sum(r.grid_export_kw for r in simulation_results)
    self_sufficiency = ((total_load - total_grid_import) / total_load * 100) if total_load > 0 else 0
    
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Light Grid Accent 1'
    
    # Header
    header_cells = table.rows[0].cells
    headers = [RL['metric'], RL['daily'], RL['monthly'], RL['annual_est']]
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(header_cells[i], 'E67E22')
    
    # Data
    data = [
        (RL['solar_generation'], f"{total_pv:.2f} kWh", f"{total_pv * 30:.1f} kWh", f"{total_pv * 365:.0f} kWh"),
        (RL['total_consumption'], f"{total_load:.2f} kWh", f"{total_load * 30:.1f} kWh", f"{total_load * 365:.0f} kWh"),
        (RL['grid_import'], f"{total_grid_import:.2f} kWh", f"{total_grid_import * 30:.1f} kWh", f"{total_grid_import * 365:.0f} kWh"),
        (RL['grid_export'], f"{total_grid_export:.2f} kWh", f"{total_grid_export * 30:.1f} kWh", f"{total_grid_export * 365:.0f} kWh"),
        (RL['self_sufficiency'], f"{self_sufficiency:.1f}%", f"{self_sufficiency:.1f}%", f"{self_sufficiency:.1f}%"),
    ]
    
    for i, row_data in enumerate(data, 1):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = cell_data
            if j == 0:
                table.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()

def add_device_inventory(doc, devices):
    """Add device inventory section"""
    doc.add_page_break()
    doc.add_heading(RL['device_inventory'], 2)
    
    # Summary
    total_device_power = sum(d.power_watts for d in devices)
    total_daily_energy = sum(d.daily_energy_kwh for d in devices)
    priority_count = sum(1 for d in devices if d.is_priority)
    
    summary = doc.add_paragraph(f"{RL['total_devices']}: {len(devices)} | {RL['total_power']}: {total_device_power:.0f}W | {RL['daily_consumption']}: {total_daily_energy:.2f} kWh | {RL['priority_devices']}: {priority_count}")
    summary.runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # Table
    table = doc.add_table(rows=len(devices) + 1, cols=7)
    table.style = 'Light Grid Accent 1'
    
    # Header
    headers = ['#', RL['device_name'], RL['power_w'], RL['hours_day'], RL['daily_energy'], RL['type'], RL['priority']]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, '8E44AD')
    
    # Data
    for idx, device in enumerate(devices, 1):
        row = table.rows[idx]
        row.cells[0].text = str(idx)
        row.cells[1].text = device.name
        row.cells[2].text = f"{device.power_watts:.0f}"
        row.cells[3].text = f"{device.daily_hours:.1f}"
        row.cells[4].text = f"{device.daily_energy_kwh:.2f} kWh"
        row.cells[5].text = device.device_type.title()
        row.cells[6].text = "⭐ Yes" if device.is_priority else "○ No"

def add_footer(doc):
    """Add professional footer"""
    doc.add_paragraph()
    doc.add_paragraph('_' * 80)
    
    footer1 = doc.add_paragraph(RL['footer_text'])
    footer1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer1.runs[0].font.size = Pt(9)
    footer1.runs[0].font.color.rgb = RGBColor(127, 140, 141)
    
    footer2 = doc.add_paragraph(RL['contact_support'])
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer2.runs[0].font.size = Pt(9)
    footer2.runs[0].font.color.rgb = RGBColor(127, 140, 141)
    
    footer3 = doc.add_paragraph(f"{RL['report_id']}: KHS-{datetime.datetime.now().strftime('%Y%m%d-%H%M%S')}")
    footer3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer3.runs[0].font.size = Pt(9)
    footer3.runs[0].font.color.rgb = RGBColor(127, 140, 141)

def generate_word_report(simulation_results: List[SimulationResult],
                        devices: List[Device],
                        financial: FinancialAnalysis,
                        system_config: dict,
                        filename: str):
    """Generate professional Word document report"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Add sections
    add_professional_header(doc)
    add_customer_info(doc, system_config)
    add_system_config(doc, system_config, len(devices))
    add_cost_breakdown(doc, system_config, financial)
    add_financial_analysis(doc, financial)
    add_energy_summary(doc, simulation_results)
    add_device_inventory(doc, devices)
    add_footer(doc)
    
    # Save document
    doc.save(filename)
